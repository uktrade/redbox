import logging
import time
from datetime import UTC, datetime
from io import BytesIO
from typing import TYPE_CHECKING, List, Tuple

import environ
import fitz
from redbox.chains.parser import ClaudeParser

from langchain_core.documents import Document
from langchain_core.prompts import PromptTemplate

from pydantic import ValidationError
from redbox_app.setting_enums import Environment

from redbox.chains.components import get_chat_llm
from redbox.models.chain import GeneratedMetadata
from redbox.models.file import TabularSchema
from redbox.models.settings import Settings
from redbox.transform import bedrock_tokeniser
import pandas as pd
import math
import boto3
from typing import Iterator

from docx import Document as DocxDocument

env = environ.Env()
ENVIRONMENT = Environment[env.str("ENVIRONMENT").upper()]

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

tokeniser = bedrock_tokeniser

if TYPE_CHECKING:
    from mypy_boto3_s3.client import S3Client
else:
    S3Client = object


def is_large_pdf(file_name: str, filebytes: BytesIO, page_threshold: int = 150) -> Tuple[bool, int]:
    if not file_name.lower().endswith(".pdf"):
        return False, 0
    try:
        doc = fitz.open(stream=filebytes.getvalue(), filetype="pdf")
        return len(doc) > page_threshold, len(doc)
    except Exception as e:
        logger.warning("error opening PDF - %s", e)
        # assume its not large if you can't open it
        return False, 0


def split_pdf(filebytes: BytesIO, pages_per_chunk: int = 75) -> List[BytesIO]:
    doc = fitz.open(stream=filebytes.getvalue(), filetype="pdf")
    chunks: List[BytesIO] = []
    total_pages = len(doc)
    if total_pages == 0:
        return chunks

    for start in range(0, total_pages, pages_per_chunk):
        sub_doc = fitz.open()
        end = min(start + pages_per_chunk, total_pages)
        sub_doc.insert_pdf(doc, from_page=start, to_page=end)
        if len(sub_doc) == 0:
            continue  # Skip empty chunks
        chunk_bytes = BytesIO(sub_doc.tobytes())
        chunks.append(chunk_bytes)
    return chunks


def _pdf_is_image_heavy(file_bytes: BytesIO, sample_pages: int = 5, image_threshold: int = 1) -> bool:
    try:
        doc = fitz.open(stream=file_bytes.getvalue(), filetype="pdf")
        pages_to_check = min(len(doc), sample_pages)
        images_found = 0
        for i in range(pages_to_check):
            page = doc[i]
            images = page.get_images(full=True)
            if images:
                images_found += 1
        # if more than half of sampled pages have images, file == image-heavy
        return images_found >= math.ceil(pages_to_check / 2)
    except Exception as e:
        logger.debug("can't work out quantity of images for the file - %s", e)
        return False


def infer_sqlite_type(dtype) -> str:
    if pd.api.types.is_integer_dtype(dtype):
        return "INTEGER"
    if pd.api.types.is_float_dtype(dtype):
        return "REAL"
    if pd.api.types.is_bool_dtype(dtype):
        return "BOOLEAN"
    return "TEXT"


def read_csv_text(file_bytes: BytesIO) -> list[dict[str, str | dict]]:
    """Reads in a csv file, validates it using pandas and then returns the csv as string with a null metadata dictionary"""
    try:
        file_bytes.seek(0)
        # Read bytes into pandas df. This acts as a pre-check that the csv is well formed
        df = pd.read_csv(file_bytes)
        if df.empty:
            logger.error("Empty File Uploaded")
            raise ValidationError("Empty File Uploaded")
        return [
            {
                "text": str(df.to_csv(index=False)),  # Convert bytes to string
                "metadata": {},  # returning empty metadata dictionary
            }
        ]
    except Exception as e:
        if isinstance(e, ValidationError):
            raise
        logger.error(f"Error while trying to upload csv file {e}")
        return None


def read_excel_file(file_bytes: BytesIO) -> list[dict[str, str | dict]]:
    """Reads in an excel file, validates each sheet using pandas and then returns a list of each valid sheet as string with a null metadata dictionary"""
    try:
        sheets = pd.read_excel(file_bytes, sheet_name=None)
        elements = []

        for name, df in sheets.items():
            try:
                if df.empty:
                    logger.info(f"Skipping Sheet {name}")
                    continue
                # Include the table name in the text that is stored. This will be extracted by the retriever
                table_name = name.lower().replace(" ", "_")
                sheet_schema = TabularSchema(
                    name=table_name, columns={col: infer_sqlite_type(df[col].dtype) for col in df.columns}
                )

                csv_text = f"<table_name>{table_name}</table_name>" + str(df.to_csv(index=False))
                elements.append({"text": csv_text, "metadata": {"document_schema": sheet_schema.model_dump()}})
            except Exception as e:
                logger.info(f"Skipping Sheet {name} due to error: {e}")
                continue
        return elements if len(elements) else None
    except Exception as e:
        logger.error(f"Excel Read Error: {e}")
        return None


def load_tabular_file(file_name: str, file_bytes: BytesIO) -> list[dict[str, str]]:
    """Selects the right read method for each file type. Returns an empty list if n"""
    if file_name.endswith(".csv"):
        elements = read_csv_text(file_bytes=file_bytes)
    else:
        elements = read_excel_file(file_bytes=file_bytes)

    return elements if elements else []


class TextractChunkLoader:
    """
    Load, partition and chunk a document using:
    - Textract for PDFs
    - python-docx for DOCX.
    - Pandas for CSV/Excel
    """

    def __init__(
        self,
        bucket: str,
        min_chunk_size: int = 500,
        max_chunk_size: int = 2000,
        overlap_chars: int = 200,
        region: str = "eu-west-2",
    ):
        self.bucket = bucket
        self.textract = boto3.client("textract", region_name=region)
        self.s3 = boto3.client("s3", region_name=region)
        self.min_chunk_size = min_chunk_size
        self.max_chunk_size = max_chunk_size
        self.overlap_chars = overlap_chars

        logger.info(
            "Initialised TextractChunkLoader (bucket=%s, region=%s, min_chunk=%s, max_chunk=%s, overlap=%s)",
            bucket,
            region,
            min_chunk_size,
            max_chunk_size,
            overlap_chars,
        )

    def _wait_for_job(self, job_id: str):
        logger.info("Waiting for Textract job %s to complete", job_id)

        while True:
            try:
                response = self.textract.get_document_text_detection(JobId=job_id)
                status = response["JobStatus"]

                logger.debug("Textract job %s current status: %s", job_id, status)

                if status in ["SUCCEEDED", "FAILED"]:
                    logger.info("Textract job %s finished with status: %s", job_id, status)
                    return status

                time.sleep(3)

            except Exception as e:
                logger.exception("Error while polling Textract job %s: %s", job_id, e)
                raise

    def _get_textract_results(self, job_id: str) -> List[str]:
        logger.info("Fetching Textract results for job %s", job_id)

        pages: dict[int, List[str]] = {}
        next_token = None
        api_calls = 0

        while True:
            try:
                kwargs = {"JobId": job_id}
                if next_token:
                    kwargs["NextToken"] = next_token

                response = self.textract.get_document_text_detection(**kwargs)
                api_calls += 1

                for block in response.get("Blocks", []):
                    if block["BlockType"] == "LINE":
                        page = block.get("Page", 1)
                        pages.setdefault(page, []).append(block["Text"])

                next_token = response.get("NextToken")
                if not next_token:
                    break

            except Exception as e:
                logger.exception("Error retrieving Textract results for job %s: %s", job_id, e)
                raise

        logger.info(
            "Retrieved Textract results for job %s: %d pages via %d API calls",
            job_id,
            len(pages),
            api_calls,
        )

        return ["\n".join(pages[p]) for p in sorted(pages)]

    def _extract_pdf_from_s3(self, bucket: str, key: str) -> list[str]:
        logger.info("Starting Textract extraction directly from S3: s3://%s/%s", bucket, key)

        try:
            response = self.textract.start_document_text_detection(
                DocumentLocation={
                    "S3Object": {
                        "Bucket": bucket,
                        "Name": key,
                    }
                }
            )

            job_id = response["JobId"]
            logger.info("Started Textract job %s for s3://%s/%s", job_id, bucket, key)

            status = self._wait_for_job(job_id)

            if status != "SUCCEEDED":
                logger.error("Textract job %s failed for s3://%s/%s", job_id, bucket, key)
                raise RuntimeError(f"Textract failed for s3://{bucket}/{key}")

            return self._get_textract_results(job_id)

        except Exception as e:
            logger.exception("Textract extraction failed for s3://%s/%s: %s", bucket, key, e)
            raise

    def _extract_docx(self, file_bytes: BytesIO) -> List[str]:
        logger.info("Extracting DOCX content")
        file_bytes.seek(0)
        doc = DocxDocument(file_bytes)
        text = "\n".join(p.text for p in doc.paragraphs)
        logger.debug("Extracted %d characters from DOCX", len(text))
        return [text] if text.strip() else []

    def _chunk_text(self, text: str) -> List[str]:
        chunks = []
        start = 0
        length = len(text)
        while start < length:
            end = min(start + self.max_chunk_size, length)
            chunk = text[start:end]
            if len(chunk) >= self.min_chunk_size:
                chunks.append(chunk)
            start = end - self.overlap_chars
        return chunks

    def lazy_load(
        self,
        file_name: str,
        s3_bucket: str,
        s3_key: str,
        file_bytes: BytesIO | None = None,
    ) -> Iterator[Document]:
        if file_name.lower().endswith((".csv", ".xls", ".xlsx")):
            if file_bytes is None:
                obj = self.s3.get_object(Bucket=s3_bucket, Key=file_name)
                file_bytes = BytesIO(obj["Body"].read())

            tabular_elements = load_tabular_file(file_name, file_bytes)
            for idx, el in enumerate(tabular_elements):
                schema_dict = el.get("metadata", {}).get("document_schema")
                schema = None
                if schema_dict:
                    try:
                        schema = TabularSchema.model_validate(schema_dict)
                    except ValidationError:
                        schema = None

                token_count = tokeniser(el["text"])
                metadata = {
                    "index": idx,
                    "uri": file_name,
                    "page_number": 1,
                    "created_datetime": datetime.now(UTC),
                    "token_count": token_count,
                    "document_schema": schema,
                }
                yield Document(page_content=el["text"], metadata=metadata)
            return

        if file_name.lower().endswith(".docx"):
            if file_bytes is None:
                obj = self.s3.get_object(Bucket=s3_bucket, Key=file_name)
                file_bytes = BytesIO(obj["Body"].read())
            pages = self._extract_docx(file_bytes)

        else:
            loader = TextractChunkLoader(bucket=s3_bucket)
            pages = loader._extract_pdf_from_s3(bucket=s3_bucket, key=s3_key)

        idx = 0
        for page_num, page_text in enumerate(pages, start=1):
            for chunk in self._chunk_text(page_text):
                metadata = {
                    "index": idx,
                    "uri": file_name,
                    "page_number": page_num,
                    "created_datetime": datetime.now(UTC),
                    "token_count": tokeniser(chunk),
                }
                yield Document(page_content=chunk, metadata=metadata)
                idx += 1


class MetadataLoader:
    """
    Extract metadata from a file using a TextractChunkLoader and LLM.
    Preserves trimming and robust handling from old loader.
    """

    def __init__(self, env: Settings, s3_client: S3Client, file_name: str):
        self.env = env
        self.s3_client = s3_client
        self.llm = get_chat_llm(env.metadata_extraction_llm)
        self.file_name = file_name

    def _get_file_bytes(self, file_name: str) -> BytesIO:
        obj = self.s3_client.get_object(Bucket=self.env.bucket_name, Key=file_name)
        return BytesIO(obj["Body"].read())

    def extract_metadata(self) -> GeneratedMetadata:
        start_time = time.time()

        loader = TextractChunkLoader(
            bucket=self.env.bucket_name,
            min_chunk_size=200,
            max_chunk_size=2000,
            overlap_chars=0,
        )

        file_bytes = None
        if self.file_name.lower().endswith(".docx"):
            file_bytes = self._get_file_bytes(self.file_name)

        try:
            chunks = list(
                loader.lazy_load(
                    file_name=self.file_name,
                    s3_bucket=self.env.bucket_name,
                    s3_key=self.file_name,
                    file_bytes=file_bytes,
                )
            )
        except Exception as e:
            logger.info("Failed metadata extraction - %s", e)
            chunks = []

        first_10k_chars = "".join(c.page_content for c in chunks)[:10_000]

        try:
            metadata = self.create_file_metadata(first_10k_chars)
        except Exception as e:
            logger.info(e)
            metadata = GeneratedMetadata(name=self.file_name)

        logger.info(
            "Total metadata extraction for file [%s] took %.2f seconds",
            self.file_name,
            time.time() - start_time,
        )

        return metadata

    def create_file_metadata(self, page_content: str, original_metadata: dict | None = None) -> GeneratedMetadata:
        """Trim original metadata and invoke LLM chain"""
        if not original_metadata:
            original_metadata = {}

        def trim(obj, max_length=1000):
            if isinstance(obj, dict):
                return {k: trim(v, max_length) for k, v in obj.items()}
            if isinstance(obj, list):
                return [trim(v, max_length) for v in obj]
            if isinstance(obj, str):
                return obj[:max_length]
            return obj

        original_metadata = trim(original_metadata)

        parser = ClaudeParser(pydantic_object=GeneratedMetadata)

        metadata_prompt = PromptTemplate(
            template="".join(self.env.metadata_prompt)
            + "\n\n{format_instructions}\n\n{page_content}\n\n{original_metadata}",
            input_variables=["page_content"],
            partial_variables={
                "format_instructions": parser.get_format_instructions(),
                "original_metadata": original_metadata,
            },
        )
        metadata_chain = metadata_prompt | self.llm | parser

        try:
            return metadata_chain.invoke({"page_content": page_content})
        except ValidationError as e:
            logger.info(e.errors())
            return GeneratedMetadata(name=original_metadata.get("filename") or self.file_name)
